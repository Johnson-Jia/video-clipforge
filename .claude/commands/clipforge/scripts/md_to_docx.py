# -*- coding: utf-8 -*-
"""md_to_docx.py — Markdown → Word 文档转换（github-zhihu 周报双格式输出用）

管线确定化：固定转换步骤代码化，LLM 不介入。pandoc 负责结构转换（标题/表格/链接/引用），
python-docx 后处理中文排版（东亚字体声明 + 表格边框），Word/WPS 打开即为成品。

用法：
  python scripts/md_to_docx.py --input article.md --output article.docx [--title 标题]

依赖：pandoc >= 3.x（anaconda 自带）+ python-docx
"""
import argparse
import subprocess
import sys
import tempfile
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

# 中文字体方案：标题黑体系 / 正文宋体系（Word 中文排版惯例）
HEADING_FONT = "微软雅黑"
BODY_FONT = "宋体"
MONO_FONT = "Consolas"
LINK_COLOR = RGBColor(0x09, 0x69, 0xDA)  # Word 默认超链接蓝


def _set_east_asian_font(run, font_name):
    """python-docx 不直接写东亚字体声明（w:eastAsia），需操作底层 rPr。"""
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = rPr.makeelement(qn("w:rFonts"), {})
        rPr.append(rFonts)
    rFonts.set(qn("w:eastAsia"), font_name)


def _style_run(run, font, size, bold=None):
    run.font.name = font
    _set_east_asian_font(run, font)
    run.font.size = Pt(size)
    if bold is not None:
        run.font.bold = bold


def _set_table_borders(table):
    """pandoc 生成的表格默认无框，补全单线边框（Word 'Table Grid' 等效）。"""
    tbl = table._tbl
    tblPr = tbl.tblPr
    borders = tblPr.makeelement(qn("w:tblBorders"), {})
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = borders.makeelement(qn(f"w:{edge}"), {})
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "4")
        el.set(qn("w:color"), "999999")
        borders.append(el)
    tblPr.append(borders)


def postprocess(docx_path, title=None):
    """pandoc 产物的中文排版后处理。"""
    doc = Document(str(docx_path))

    # 标题体系：黑体系，字号随层级递减
    for para in doc.paragraphs:
        style_name = para.style.name if para.style else ""
        if style_name.startswith("Heading"):
            for run in para.runs:
                level = style_name.replace("Heading ", "")
                size = {"1": 20, "2": 16, "3": 14}.get(level, 12)
                _style_run(run, HEADING_FONT, size, bold=True)
        elif style_name.startswith("Title"):
            for run in para.runs:
                _style_run(run, HEADING_FONT, 22, bold=True)
        else:
            for run in para.runs:
                _style_run(run, BODY_FONT, 11)

    # 超链接着色（pandoc 转换后 link run 可能是默认黑色，统一为链接蓝）
    for para in doc.paragraphs:
        for run in para.runs:
            if run._element.findall(qn("w:hyperlink")) or "Hyperlink" in (
                run.style.name if run.style else ""
            ):
                run.font.color.rgb = LINK_COLOR

    # 表格：边框 + 表头加粗 + 中文字体
    for table in doc.tables:
        _set_table_borders(table)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        for row_idx, row in enumerate(table.rows):
            for cell in row.cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        _style_run(run, HEADING_FONT if row_idx == 0 else BODY_FONT,
                                   10, bold=(row_idx == 0))

    doc.save(str(docx_path))


def convert(md_path, docx_path, title=None):
    md_path, docx_path = Path(md_path), Path(docx_path)
    if not md_path.exists():
        print(f"[md_to_docx] 输入不存在: {md_path}", file=sys.stderr)
        return 1

    # pandoc 结构转换（临时文件放目标同目录——Windows 的 os.replace 不能跨盘）
    tmp_out = docx_path.parent / f".{docx_path.stem}.tmp{docx_path.suffix}"
    try:
        cmd = ["pandoc", str(md_path), "-f", "markdown", "-t", "docx",
               "--toc-depth=2", "-o", str(tmp_out)]
        if title:
            cmd += ["--metadata", f"title={title}"]
        subprocess.run(cmd, check=True, capture_output=True)
        postprocess(tmp_out, title)
        docx_path.parent.mkdir(parents=True, exist_ok=True)
        tmp_out.replace(docx_path)
    finally:
        tmp_out.unlink(missing_ok=True)

    print(f"[md_to_docx] 已生成: {docx_path}（{docx_path.stat().st_size:,} bytes）")
    return 0


if __name__ == "__main__":
    ap = argparse.ArgumentParser(description="Markdown → docx（github-zhihu 双格式输出）")
    ap.add_argument("--input", required=True, help="article.md 路径")
    ap.add_argument("--output", default=None, help="输出 docx 路径（默认同目录同名 .docx）")
    ap.add_argument("--title", default=None, help="文档元标题（可选）")
    args = ap.parse_args()
    out = args.output or str(Path(args.input).with_suffix(".docx"))
    sys.exit(convert(args.input, out, args.title))
